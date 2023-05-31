# encoding="utf-8"
#
# This function takes the google sheet list of publications/abstracts from either the google drive website or from
# the downloaded csv file, and output the list of publications (or abstracts or talks) in either format: word, ccv,
# txt for neuropoly wiki.
# Developed by: Alexandru Foias and Julien Cohen-Adad

import csv
import argparse
import os
import re
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import pprint
import pandas as pd

from docx import Document
from docx.shared import Inches


def get_parameters():
    parser = argparse.ArgumentParser(description='This script is used for'
        'generating the publications content for neurowiki')
    parser.add_argument("-j", "--json_auth",
                        help="Path to csv containing publications",
                        required=True)
    args = parser.parse_args()
    return args


def main(path_json_auth):
    """
    Main function
    :param path_json_auth:
    :return:
    """

    # Populate pandas dataframe from either a downloaded csv file, or by connecting to Google sheet via a json
    # authentification file
    use_csv = True
    if use_csv:
        df_publis = pd.read_csv(path_json_auth, delimiter=',')

    else:
        scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
        creds = ServiceAccountCredentials.from_json_keyfile_name(path_json_auth, scope)
        client = gspread.authorize(creds)

        sheet = client.open ('publications').sheet1

        # TODO: convert to Pandas dataframe
        df_publis = sheet.get_all_records()

        list_csv_title = []
        list_csv_author = []
        list_csv_journal_name = []
        list_csv_journal_info = []

    # TODO: Make it a flag when calling the function
    type_output = 'word'

    # Loop across df lines and output formatted text file
    for hh in range(0, len(df_publis)):
        print(hh)
        authors = df_publis['Authors'][hh]
        title = df_publis['Title'][hh]
        conference = df_publis['Conference'][hh]
        date = str(df_publis['Date'][hh])

        #
        # csv_journal = df_publis[hh]['Reference'].split('. ')[2:]
        # csv_journal = ' '.join(csv_journal)
        # csv_journal_name = re.split(r'(\d+)',csv_journal) #strip journal name assuming that the name ends at the first numeric value
        # if len(csv_journal_name) > 1: # grab journal information year,doi etc.
        #     csv_journal_info = csv_journal.split(csv_journal_name[0])
        #     csv_journal_info = ''.join(csv_journal_info)
        # else:
        #     csv_journal_info = '' # generate void content if journal information year,doi etc. doesn't exist
        #
        # list_csv_title.append(csv_title)
        # list_csv_author.append(csv_author)
        # list_csv_journal_name.append(csv_journal_name[0])
        # list_csv_journal_info.append(csv_journal_info)

        if type_output == 'word':

            # Initialize word document
            document = Document()

            p = document.add_paragraph(authors, style='List Number')
            p.add_run('. ')
            p.add_run(title).italic = True
            p.add_run('. ')
            p.add_run(conference).bold = True
            p.add_run('. ')
            p.add_run(date)

        # p = document.add_paragraph('A plain paragraph having some ')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True
    document.save('demo.docx')



    list_output_txt = [] #generate text format for https://www.neuro.polymtl.ca/publications#journal_articles respecting dokuwiki formatting parameters
    for row_id in reversed (range(len(list_csv_author))):
        list_output_txt.append('   - ' +  list_csv_author[row_id] + '. [[' + list_csv_title[row_id] + '|' + list_csv_title[row_id] + ']] ' + '**' + list_csv_journal_name[row_id] + '** ' + list_csv_journal_info[row_id])

    buffer_root_path_output_csv = path_json_auth.split('/') # generate path for output file; saved in the same location as the csv file


    root_path_output_csv = '/'.join(str(buffer_root_path_output_csv[i]) for i in range (len(buffer_root_path_output_csv)-1))

    for item in list_output_txt:
        print(item)
    # with open(root_path_output_csv + '/content_publications_neurowiki.txt', 'w') as txtFile: #write new generated publications in txt file
    #     for item in list_output_txt:
    #         txtFile.write("%s\n" % item)
    # txtFile.close()

if __name__ == "__main__":
    args = get_parameters()
    main(args.json_auth)
