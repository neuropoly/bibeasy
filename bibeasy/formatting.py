#!/usr/bin/env python
#
# Deals with output formatting

import logging
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd


# TODO: fetch students from google sheet
STUDENTS = [
    'Alley S',
    'Alonso-Ortiz E',
    'Badji A',
    'Benhamou M',
    'Boudreau M',
    'Bourget M-H',
    'De Leener B',
    'Dupont S',
    'Dupont SM',
    'Duval T',
    'Eden D',
    'Enguix V',
    'Foias A',
    'Germain G',
    'Gros C',
    'Kerbrat A',
    'Lemay A',
    'Levy S',
    'Lopez Rios N',
    'Lubrano M',
    'Lubrano di Scandalea M',
    'Mangeat G',
    'Mingasson T',
    'Morozov D',
    'Nami H',
    'Paquin ME',
    'Paugam F',
    'Perdigon Romero F',
    'Perone C',
    'Perone CS',
    'Perraud B',
    'Rouhier L',
    'Saliani A',
    'Snoussi H',
    'Topfer R',
    'Ullman E',
    'Verma T',
    'Vincent O',
    'Wabartha M',
    'Zaimi A',
    ]


def check_field_exists(field):
    def inner(func):
        def wrapper(paragraph, row):
            if field in row:
                if row[field]:
                    func(paragraph, row)
        return wrapper
    return inner


def convert_labels_file(labels_file, output_file):
    """
    Convert a text file with labels to a HTML file with buttons. Used for NeuroPoly's website.
    """
    # Read the labels from the input file
    with open(labels_file, "r") as f:
        labels = [line.strip() for line in f]

    # Create the HTML code with the labels
    html_code = "<!-- label_definitions.md -->\n\n"
    for label in labels:
        html_code += f'<button class="label" data-label="{label}">{label}</button>\n'

    # Write the HTML code to the output file
    with open(output_file, "w") as f:
        f.write(html_code)


def csv_to_txt_pubtype(df, pubtype, args):
    """
    Write formatted output file with list of publication for a specific pubtype.
    :param df: DF of a single pubtype
    :param pubtype: str: Publication type
    :param args:
    :return:
    """
    file_output, ext_output = args.output.split('.')
    fname_output = '{}-{}.{}'.format(file_output, pubtype, ext_output)

    # For Dokuwiki
    if ext_output == 'md':
        # Generate aggregated list
        list_output_txt = []
        unique_years = sorted(df['Year'].unique(), reverse=True)

        for year in unique_years:
            list_output_txt.append('\n## {}'.format(year))
            # This is needed for formatting the publications in the website
            list_output_txt.append(f'<div class="publications-container">')
            df_year = df[df['Year'] == year]

            for index, row in df_year.iterrows():
                logging.debug(row)
                list_output_txt.append(_format_website(row))

        list_output_txt.append(f"</div>")

        # Save to file
        with open(fname_output, 'w') as txtFile:
            for item in list_output_txt:
                txtFile.write("%s\n" % item)
        txtFile.close()

        # Create authorized labels file
        convert_labels_file(args.labels,
                            "labels_publication.html")

    # For Word/GoogleDoc
    elif ext_output == 'docx':
        mydoc = docx.Document()
        for index, row in df.iterrows():
            logging.debug(row)
            _format_docx(mydoc, row, args.style)
        mydoc.save(fname_output)

    logging.info("\nFormatting type: '{}'".format(pubtype))
    logging.info("  Selected entries: {}".format(len(df)))
    logging.info('  File written: {}'.format(fname_output))


def _format_docx(mydoc, row, style):
    """
    Format DF row according to word/googledoc format
    :param mydoc: docx.mydoc object
    :param row: DF row
    :param style: {APA, custom}
    :return:
    """

    def _add_formatted(paragraph, txt, style=None, color='black'):
        """
        Add special formatting to docx object.
        :param paragraph: docx.Document().add_paragraph() object
        :param txt:
        :param style:
        :param color:
        :return:
        """
        color_to_rgbstr = {
            'black': '000000',
            'blue': '0000FF',
            'red': 'FF0000',
            'green': '008000'
            }
        run = paragraph.add_run(txt)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        if style == 'bold':
            font.bold = True
        elif style == 'italic':
            font.italic = True
        elif style == 'underline':
            font.underline = True
        font.color.rgb = RGBColor.from_string(color_to_rgbstr[color])

    @check_field_exists('Authors')
    def _add_authors(paragraph, row):
        """Format list of authors for based on special characteristics, e.g. if they are students or HQP."""
        list_authors = [a.strip() for a in row['Authors'].split(',')]
        for author in list_authors:
            # If student, apply special formatting
            if author in STUDENTS:
                style = 'underline'
            else:
                style = None
            _add_formatted(paragraph, "{}".format(author), style=style)
            # Add "," or "." depending if last author
            if not author == list_authors[-1]:
                _add_formatted(paragraph, ", ")

    @check_field_exists('Impact')
    def _add_impact(paragraph, row):
        _add_formatted(paragraph, " (IF: {})".format(row['Impact']))

    @check_field_exists('Location')
    def _add_location(paragraph, row):
        _add_formatted(paragraph, ", ({})".format(row['Location']))

    @check_field_exists('Prize')
    def _add_prize(paragraph, row):
        _add_formatted(paragraph, ". {}".format(row['Prize']), style='italic', color='blue')

    @check_field_exists('URL')
    def _add_url(paragraph, row):
        # Only add for pubtype=media
        # if row['Type'] == 'media':
        _add_formatted(paragraph, ". {}".format(row['URL']), style='italic', color='blue')

    @check_field_exists('Volume:Pages')
    def _add_volume(paragraph, row):
        _add_formatted(paragraph, ", {}".format(row['Volume:Pages']))

    paragraph = mydoc.add_paragraph()
    # Custom value to match Julien's docx publication file
    paragraph.paragraph_format.left_indent = Inches(0.58)
    paragraph.paragraph_format.first_line_indent = Inches(-0.58)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if style == 'APA':
        _add_authors(paragraph, row)
        _add_formatted(paragraph, ". ({})".format(row['Year']))
        _add_formatted(paragraph, ". {}".format(row['Title']))
        _add_formatted(paragraph, ". {}".format(row['Journal/Conference']), style='italic')
        _add_location(paragraph, row)
        _add_impact(paragraph, row)
        _add_volume(paragraph, row)
        _add_prize(paragraph, row)
        _add_url(paragraph, row)

    elif style == 'custom':
        _add_formatted(paragraph, "[{}]\t".format(row['ID']), style='bold')
        _add_authors(paragraph, row)
        # Check special case for talks (no authors)
        if 'Authors' in row:
            if not row['Authors'] == '':
                prefix_auth = ". "
            else:
                prefix_auth = ""
        else:
            prefix_auth = ""
        _add_formatted(paragraph, "{}{}".format(prefix_auth, row['Title']), style='italic')
        _add_formatted(paragraph, ". ")  # TODO: merge this line with the one below, and incorporate the ". " directly in the _add_formatted functions
        _add_formatted(paragraph, "{}".format(row['Journal/Conference']), style='bold')
        _add_location(paragraph, row)
        _add_impact(paragraph, row)
        _add_volume(paragraph, row)
        _add_formatted(paragraph, ", {}".format(row['Year']))
        _add_prize(paragraph, row)
        # _add_url(paragraph, row)

    elif style == 'talk':
        _add_formatted(paragraph, "{}".format(row['Title']))
        _add_formatted(paragraph, ". {}".format(row['Journal/Conference']), style='italic')
        _add_location(paragraph, row)
        _add_impact(paragraph, row)
        _add_volume(paragraph, row)
        _add_formatted(paragraph, ", {}. ".format(row['Year']))
        _add_prize(paragraph, row)
        _add_url(paragraph, row)

    else:
        raise ValueError("You need to select a style.")


def _format_website(row):
    """
    Format DF row for NeuroPoly's website.
    https://www.neuro.polymtl.ca/publications
    :param row: DF row
    :return: str: Formatted row
    """
    labels = row['Labels'] if not pd.isnull(row['Labels']) else ''  # Check if 'Labels' is empty
    data_labels = ' '.join(labels.split(', '))  # Converts "Label 1, Label 2" to "label1 label2"
    label_info = f' (Labels: {labels})' if labels else ''  # Add label info only if labels are not empty

    formatted_row = (
        '<div class="publication" data-labels="{}">\n'
        '    <h3>{}</h3>\n'
        '    <p><em>{}</em></p>\n'
        '    <p><strong>{}</strong> ({}) <a href="{}">Link to paper</a><span class="publication-label"> (Labels: {})</span></p>\n'
        '</div>'
    ).format(
        data_labels,
        row['Title'],
        row['Authors'],
        row['Journal/Conference'],
        row['Year'],
        row['URL'],
        labels
    )

    return formatted_row
